import csv
from datetime import date
from io import BytesIO, StringIO
from typing import Annotated, Literal

from docx import Document
from fastapi import APIRouter, HTTPException, Query, Response
from openpyxl import Workbook
from openpyxl.styles import Font
from reportlab.lib.pagesizes import letter
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet

from app.api.deps import CurrentUserDep, DbDep
from app.services.range_summary import RangeReport, build_range_report
from app.services.recurring import materialize_due

router = APIRouter(prefix="/reports", tags=["reports"])

ExportFormat = Literal["csv", "xlsx", "docx", "pdf"]
_MIME_TYPES = {
    "csv": "text/csv; charset=utf-8",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}


def _summary_rows(report: RangeReport) -> list[list[str | float]]:
    return [[row.month, row.income, row.expense, row.net] for row in report.monthly]


def _total_rows(report: RangeReport) -> list[list[float]]:
    income = sum(row.income for row in report.monthly)
    expense = sum(row.expense for row in report.monthly)
    return [[round(income, 2), round(expense, 2), round(income - expense, 2)]]


def _category_rows(report: RangeReport) -> list[list[str | float]]:
    return [
        [report.category_names.get(row.category_id, row.category_id), row.total]
        for row in report.by_category
    ]


def _transaction_rows(report: RangeReport) -> list[list[str | float]]:
    return [
        [
            row.date.isoformat(), row.type, row.amount, row.category, row.account,
            row.member, row.note or "",
        ]
        for row in report.transactions
    ]


def _csv_export(report: RangeReport) -> bytes:
    output = StringIO(newline="")
    writer = csv.writer(output)
    writer.writerow(["Resumen"])
    writer.writerow(["Ingresos", "Gastos", "Neto"])
    writer.writerows(_total_rows(report))
    writer.writerow([])
    writer.writerow(["Totales mensuales"])
    writer.writerow(["Mes", "Ingresos", "Gastos", "Neto"])
    writer.writerows(_summary_rows(report))
    writer.writerow([])
    writer.writerow(["Gastos por categoría"])
    writer.writerow(["Categoría", "Total"])
    writer.writerows(_category_rows(report))
    writer.writerow([])
    writer.writerow(["Movimientos detallados"])
    writer.writerow(["Fecha", "Tipo", "Monto", "Categoría", "Cuenta", "Miembro", "Nota"])
    writer.writerows(_transaction_rows(report))
    return output.getvalue().encode("utf-8-sig")


def _xlsx_export(report: RangeReport) -> bytes:
    workbook = Workbook()
    summary = workbook.active
    summary.title = "Summary"
    summary.append(["Resumen"])
    summary.append(["Ingresos", "Gastos", "Neto"])
    for row in _total_rows(report):
        summary.append(row)
    summary.append([])
    summary.append(["Totales mensuales"])
    summary.append(["Mes", "Ingresos", "Gastos", "Neto"])
    for row in _summary_rows(report):
        summary.append(row)
    summary.append([])
    summary.append(["Gastos por categoría"])
    summary.append(["Categoría", "Total"])
    for row in _category_rows(report):
        summary.append(row)
    for cell in summary[1] + summary[2] + summary[5] + summary[6] + summary[8 + len(report.monthly)] + summary[9 + len(report.monthly)]:
        cell.font = Font(bold=True)
    summary.freeze_panes = "A3"
    for column in summary.columns:
        summary.column_dimensions[column[0].column_letter].width = max(
            len(str(cell.value or "")) for cell in column
        ) + 2

    transactions = workbook.create_sheet("Transactions")
    transactions.append(["Fecha", "Tipo", "Monto", "Categoría", "Cuenta", "Miembro", "Nota"])
    for row in _transaction_rows(report):
        transactions.append(row)
    for cell in transactions[1]:
        cell.font = Font(bold=True)
    transactions.freeze_panes = "A2"
    for column in transactions.columns:
        transactions.column_dimensions[column[0].column_letter].width = min(
            max(len(str(cell.value or "")) for cell in column) + 2, 40
        )

    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def _docx_export(report: RangeReport, from_date: date, to_date: date) -> bytes:
    document = Document()
    document.add_heading("Reporte financiero", 0)
    document.add_paragraph(f"Periodo: {from_date.isoformat()} a {to_date.isoformat()}")
    _docx_table(document, "Resumen", ["Ingresos", "Gastos", "Neto"], _total_rows(report))
    _docx_table(document, "Totales mensuales", ["Mes", "Ingresos", "Gastos", "Neto"], _summary_rows(report))
    _docx_table(document, "Gastos por categoría", ["Categoría", "Total"], _category_rows(report))
    _docx_table(document, "Movimientos detallados", ["Fecha", "Tipo", "Monto", "Categoría", "Cuenta", "Miembro", "Nota"], _transaction_rows(report))
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _docx_table(document: Document, title: str, headers: list[str], rows: list[list[str | float]]) -> None:
    document.add_heading(title, level=1)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Shading Accent 1"
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        cell.text = header
    for row in rows:
        for cell, value in zip(table.add_row().cells, row, strict=True):
            cell.text = str(value)


def _pdf_export(report: RangeReport, from_date: date, to_date: date) -> bytes:
    output = BytesIO()
    styles = getSampleStyleSheet()
    story = [
        Paragraph("Reporte financiero", styles["Title"]),
        Paragraph(f"Periodo: {from_date.isoformat()} a {to_date.isoformat()}", styles["Normal"]),
        Spacer(1, 12),
    ]
    _pdf_table(story, "Resumen", ["Ingresos", "Gastos", "Neto"], _total_rows(report), styles)
    _pdf_table(story, "Totales mensuales", ["Mes", "Ingresos", "Gastos", "Neto"], _summary_rows(report), styles)
    _pdf_table(story, "Gastos por categoría", ["Categoría", "Total"], _category_rows(report), styles)
    _pdf_table(story, "Movimientos detallados", ["Fecha", "Tipo", "Monto", "Categoría", "Cuenta", "Miembro", "Nota"], _transaction_rows(report), styles)
    SimpleDocTemplate(output, pagesize=letter).build(story)
    return output.getvalue()


def _pdf_table(story: list, title: str, headers: list[str], rows: list[list[str | float]], styles) -> None:
    story.extend([Paragraph(title, styles["Heading2"]), Spacer(1, 6)])
    table = Table([headers, *[[str(value) for value in row] for row in rows]], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f4e78")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.extend([table, Spacer(1, 12)])


@router.get("/export")
def export_report(
    db: DbDep,
    user: CurrentUserDep,
    format: ExportFormat,
    from_date: Annotated[date, Query(alias="from")],
    to_date: Annotated[date, Query(alias="to")],
) -> Response:
    if user.household_id is None:
        raise HTTPException(status_code=400, detail="El usuario no pertenece a un hogar")
    if from_date > to_date:
        raise HTTPException(status_code=422, detail="La fecha inicial debe ser anterior a la final")

    materialize_due(db, user.household_id, user.id)
    report = build_range_report(db, user.household_id, from_date, to_date)
    renderers = {"csv": _csv_export, "xlsx": _xlsx_export}
    if format in renderers:
        content = renderers[format](report)
    elif format == "docx":
        content = _docx_export(report, from_date, to_date)
    else:
        content = _pdf_export(report, from_date, to_date)
    filename = f"reporte-{from_date.isoformat()}-a-{to_date.isoformat()}.{format}"
    return Response(
        content=content,
        media_type=_MIME_TYPES[format],
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
